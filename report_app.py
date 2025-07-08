import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
import os
import json
import sys
import re
import openpyxl
import tkinter.font as tkFont
from tkinter import ttk
from docx import Document
from openpyxl import load_workbook
from tkinterdnd2 import DND_FILES, TkinterDnD
from idlelib.tooltip import Hovertip


class Tooltip:
    def __init__(self, widget, text, delay=500):
        self.widget = widget
        self.text = text
        self.delay = delay
        self.tip_window = None
        self.id = None
        self.x = self.y = 0

        widget.bind("<Enter>", self.schedule)
        widget.bind("<Leave>", self.unschedule)
        widget.bind("<Motion>", self.track_mouse)

    def schedule(self, event=None):
        self.unschedule()
        self.id = self.widget.after(self.delay, self.show_tip)

    def unschedule(self, event=None):
        if self.id:
            self.widget.after_cancel(self.id)
            self.id = None
        self.hide_tip()

    def track_mouse(self, event):
        self.x = event.x_root + 10
        self.y = event.y_root + 10
        if not self.tip_window:
            self.schedule()

    def show_tip(self):
        if self.tip_window or not self.text:
            return
        self.tip_window = tw = tk.Toplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{self.x}+{self.y}")
        label = tk.Label(tw, text=self.text, justify="left",
                         background="#ffffe0", relief="solid", borderwidth=1,
                         font=("Segoe UI", 9))
        label.pack(ipadx=5, ipady=2)

    def hide_tip(self):
        if self.tip_window:
            self.tip_window.destroy()
            self.tip_window = None

class TemplateFillerApp:
    def __init__(self, root):
        self.root = root
        self.root.drop_target_register(DND_FILES)
        self.root.dnd_bind("<<Drop>>", self.handle_drop)
        self.root.title("Template Filler (Word + Excel files)")
        self.placeholders = []
        self.template_type = None

        controls_frame = tk.Frame(root)
        controls_frame.grid(row=5, column=0, columnspan=4, pady=(10, 5))
        
        template_frame = tk.Frame(root)
        template_frame.grid(row=0, column=0, columnspan=4, pady=(10, 5))
        self.template_path = tk.StringVar()
        tk.Label(template_frame, text="Template File (.docx or .xlsx):").grid(row=0, column=0, sticky="w", padx=(10, 5), pady=(10, 5))
        tk.Entry(template_frame, textvariable=self.template_path, width=50).grid(row=0, column=1, padx=5, pady=(10, 5))
        ttk.Button(template_frame, text="Browse", command=self.browse_template).grid(row=0, column=2, padx=(5, 15), pady=(10, 5))

        # Create a canvas and scrollbar
        style = ttk.Style()
        style.theme_use("default")  # Or "clam" for better styling support

        style.configure(
            "Vertical.TScrollbar",
            gripcount=0,
            background="#888",
            darkcolor="#666",
            lightcolor="#aaa",
            troughcolor="#f0f0f0",
            bordercolor="#ccc",
            arrowcolor="#333",
            width=16  # Thicker scrollbar
        )

        self.scroll_container = tk.Frame(self.root)
        self.scroll_container.grid(row=1, column=0, columnspan=4, sticky="nsew")
        self.canvas = tk.Canvas(self.scroll_container, highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(
            self.scroll_container,
            orient="vertical",
            command=self.canvas.yview,
            style="Vertical.TScrollbar"
        )
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        self.canvas.grid(row=0, column=0, sticky="nsew")
        self.scrollbar.grid(row=0, column=1, sticky="ns")

        # Make canvas expand
        self.scroll_container.grid_rowconfigure(0, weight=1)
        self.scroll_container.grid_columnconfigure(0, weight=1)
        # Create a frame inside the canvas
        self.entries_frame = tk.Frame(self.canvas)
        self.entries_frame.grid_columnconfigure(0, weight=1)  # Placeholder name
        self.entries_frame.grid_columnconfigure(1, weight=1)  # Value
        self.entries_frame.grid_columnconfigure(2, weight=0)  # Delete button

        self.canvas.create_window((0, 0), window=self.entries_frame, anchor="nw")
        
        # Update scrollregion when the frame changes
        self.entries_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.enable_mousewheel_scrolling()
        self.add_table_headers()
        self.placeholder_frame = tk.Frame(root)
        self.placeholder_frame.grid(row=1, column=0, columnspan=3, padx=10, pady=10)
        self.add_placeholder_row()
        
        button_frame = tk.Frame(root)
        button_frame.grid(row=3, column=0, columnspan=4, pady=(5, 10))

        btn_add = tk.Button(button_frame, text="➕ Add Row", command=self.add_placeholder_row)
        btn_add.grid(row=0, column=0, padx=5)
        Tooltip(btn_add, "Add a new placeholder row")

        btn_clear = tk.Button(button_frame, text="🧼 Clear All Fields", command=self.clear_all_fields)
        btn_clear.grid(row=0, column=1, padx=5)
        Tooltip(btn_clear, "Clear all placeholder entries")

        btn_generate = tk.Button(button_frame, text="📄 Generate Output", command=self.generate_output)
        btn_generate.grid(row=0, column=2, padx=5)
        Tooltip(btn_generate, "Generate the final document using the template")

        btn_save = tk.Button(button_frame, text="💾 Save Profile", command=self.save_placeholders)
        btn_save.grid(row=0, column=3, padx=5)
        Tooltip(btn_save, "Save current placeholders as a profile")

        btn_delete = tk.Button(button_frame, text="🧹 Delete Profile", command=self.delete_selected_profile)
        btn_delete.grid(row=0, column=4, padx=5)
        Tooltip(btn_delete, "Delete the selected profile from disk")

        profile_frame = tk.Frame(root)
        profile_frame.grid(row=2, column=0, columnspan=4, pady=(5, 10))

        tk.Label(controls_frame, text="Profiles:").grid(row=0, column=0, sticky="e", padx=(5, 5), pady=5)

        self.profile_var = tk.StringVar()
        self.profile_var.set("Select Profile")
        self.profile_menu = ttk.Combobox(controls_frame, textvariable=self.profile_var, state="readonly", width=30)
        self.profile_menu.grid(row=0, column=1, columnspan=4, sticky="w", padx=(0, 15), pady=5)

        self.profile_menu.bind("<<ComboboxSelected>>", lambda e: self.load_named_profile())
        self.refresh_profile_menu()
        instructions = (
            "🛠️ How to Use:\n"
            "1. Select or Drag & Drop a .docx or .xlsx template with placeholders like {name}, {date}, etc.\n"
            "2. Program will attempt to auto load the placeholders into the table\n"
            "3. Add fields below to match those placeholders with values.\n"
            "4. Click 'Generate Output' to create a filled document.\n"
            "5. Use 'Save Profile' and 'Load Profile' to reuse placeholder sets anytime.\n"

        )
        tk.Label(root, text=instructions, justify="left", wraplength=800, fg="gray").grid(row=6, column=0, columnspan=4, pady=(10, 0))

    def browse_template(self):
        file_path = filedialog.askopenfilename(
            filetypes=[
                ("All Supported Files", "*.docx *.xlsx")
            ]
        )
        if file_path:
            self.template_path.set(file_path)

            # Auto-suggest placeholders
            try:
                keys = self.extract_placeholders_from_template(file_path)
                if keys:
                    self.clear_all_fields()

                    for i, key in enumerate(keys):
                        if i < len(self.placeholders):
                            self.placeholders[i][0].set(key)
                            self.placeholders[i][1].set("")  # Clear value
                        else:
                            self.add_placeholder_row()
                            self.placeholders[-1][0].set(key)
                            self.placeholders[-1][1].set("")  # Clear value
                else:
                    messagebox.showinfo("No Placeholders Found", "No placeholders were detected in the template.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to extract placeholders:\n{e}")

    def add_table_headers(self):
        tk.Label(self.entries_frame, text="Placeholder Name", font=("Segoe UI", 10, "bold")).grid(row=0, column=0, padx=5, pady=(0, 5))
        tk.Label(self.entries_frame, text="Value to Insert", font=("Segoe UI", 10, "bold")).grid(row=0, column=1, padx=5, pady=(0, 5))

    def add_placeholder_row(self, key_text="", value_text=""):
        row = len(self.placeholders) + 1  # +1 to account for header row

        key_var = tk.StringVar(value=key_text)
        val_var = tk.StringVar(value=value_text)

        key_entry = tk.Entry(self.entries_frame, textvariable=key_var, width=30)
        val_entry = tk.Entry(self.entries_frame, textvariable=val_var, width=30)

        delete_btn = tk.Button(
            self.entries_frame,
            text="❌",
            command=lambda: self.delete_placeholder_row(row - 1),
            width=3,
            font=("Segoe UI", 10)
        )


        duplicate_btn = tk.Button(
            self.entries_frame,
            text="🔁",
            command=lambda: self.duplicate_placeholder_row(row - 1),
            width=3,
            font=("Segoe UI", 10)
        )

        key_entry.grid(row=row, column=0, padx=5, pady=2, sticky="ew")
        val_entry.grid(row=row, column=1, padx=5, pady=2, sticky="ew")
        delete_btn.grid(row=row, column=2, padx=2, pady=2)
        duplicate_btn.grid(row=row, column=3, padx=2, pady=2)

        self.placeholders.append((key_var, val_var, key_entry, val_entry, delete_btn, duplicate_btn))
        self.update_delete_buttons()
        self.update_scrollbar_visibility()
        self.update_canvas_height()
        self.enable_mousewheel_scrolling()
    
    def update_delete_buttons(self):
        for i, (_, _, _, _, delete_btn, _) in enumerate(self.placeholders):
            if len(self.placeholders) == 1:
                delete_btn.config(state="disabled")
            else:
                delete_btn.config(state="normal")
    
    def delete_placeholder_row(self, index):
        _, _, key_entry, val_entry, delete_btn, duplicate_btn = self.placeholders[index]
        key_entry.destroy()
        val_entry.destroy()
        delete_btn.destroy()
        duplicate_btn.destroy()


        del self.placeholders[index]

        for i, (key_var, val_var, key_entry, val_entry, delete_btn, duplicate_btn) in enumerate(self.placeholders):
            row = i + 1
            key_entry.grid(row=row, column=0, padx=5, pady=2, sticky="ew")
            val_entry.grid(row=row, column=1, padx=5, pady=2, sticky="ew")
            delete_btn.grid(row=row, column=2, padx=2, pady=2)
            duplicate_btn.grid(row=row, column=3, padx=2, pady=2)

            delete_btn.config(command=lambda idx=i: self.delete_placeholder_row(idx))
            duplicate_btn.config(command=lambda idx=i: self.duplicate_placeholder_row(idx))

            self.update_delete_buttons()
            self.update_scrollbar_visibility()
            self.update_canvas_height()

    def clear_all_fields(self):
        for widget in self.entries_frame.winfo_children():
            widget.destroy()
        self.placeholders.clear()
        self.add_table_headers()
        self.add_placeholder_row()
    
    def update_scrollbar_visibility(self):
        self.root.update_idletasks()
        needs_scroll = self.canvas.bbox("all")[3] > self.canvas.winfo_height()
        if needs_scroll:
            self.scrollbar.grid()
        else:
            self.scrollbar.grid_remove()
    
    def enable_mousewheel_scrolling(self):
        def _on_mousewheel(event):
            if sys.platform == 'darwin':
                self.canvas.yview_scroll(-1 * int(event.delta), "units")
            else:
                self.canvas.yview_scroll(-1 * int(event.delta / 120), "units")

        def bind_mousewheel(widget):
            widget.bind("<Enter>", lambda e: self.canvas.bind_all("<MouseWheel>", _on_mousewheel))
            widget.bind("<Leave>", lambda e: self.canvas.unbind_all("<MouseWheel>"))

        # Bind to the scrollable frame itself
        bind_mousewheel(self.entries_frame)

        # Also bind to all existing widgets inside it
        for child in self.entries_frame.winfo_children():
            bind_mousewheel(child)
    
    def update_canvas_height(self):
        self.root.update_idletasks()
        content_height = self.entries_frame.winfo_reqheight()
        max_height = 200 # You can adjust this

        if content_height > max_height:
            self.canvas.config(height=max_height)
            self.scrollbar.grid()
        else:
            self.canvas.config(height=content_height)
            self.scrollbar.grid_remove()

    def generate_output(self):
        try:
            path = self.template_path.get()
            if not path or not os.path.exists(path):
                messagebox.showerror("Missing Template", "Please select a valid .docx or .xlsx template.")
                return

            if not (path.lower().endswith(".docx") or path.lower().endswith(".xlsx")):
                messagebox.showerror("Unsupported Format", "Only .docx and .xlsx templates are supported.")
                return

            # Build placeholder dictionary
            replacements = {}
            for key_var, val_var, *_ in self.placeholders:
                key = key_var.get().strip()
                val = val_var.get().strip()
                if key:
                    replacements[key] = val

            # Set default extension and filetypes BEFORE the dialog
            if path.lower().endswith(".docx"):
                default_ext = ".docx"
                filetypes = [("Word Document", "*.docx")]
            else:
                default_ext = ".xlsx"
                filetypes = [("Excel Workbook", "*.xlsx")]

            # Ask user where to save the output
            output_path = filedialog.asksaveasfilename(
                defaultextension=default_ext,
                filetypes=filetypes,
                title="Save Filled Document As"
            )
            if not output_path:
                return  # User cancelled

            # Generate output based on file type
            if path.lower().endswith(".docx"):
                self.generate_word_output(path, output_path, replacements)
            elif path.lower().endswith(".xlsx"):
                self.generate_excel_output(path, output_path, replacements)

            messagebox.showinfo("Success", f"Output saved to:\n{output_path}")

        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate output:\n{e}")

    def generate_word_output(self, template_path, output_path, replacements):
        doc = Document(template_path)
        for para in doc.paragraphs:
            for key, val in replacements.items():
                para.text = para.text.replace(f"{{{key}}}", val)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, val in replacements.items():
                        cell.text = cell.text.replace(f"{{{key}}}", val)

        doc.save(output_path)

    def generate_excel_output(self, template_path, output_path, replacements):
        wb = load_workbook(template_path)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        for key, val in replacements.items():
                            cell.value = cell.value.replace(f"{{{key}}}", val)
        wb.save(output_path)


    def get_profiles_dir(self):
        base = os.path.dirname(sys.executable if getattr(sys, 'frozen', False) else os.path.abspath(__file__))
        path = os.path.join(base, "profiles")
        os.makedirs(path, exist_ok=True)
        return path

    def save_placeholders(self):
        # Extract key-value pairs from the placeholder rows
        data = {
            key_var.get(): val_var.get()
            for key_var, val_var, *_ in self.placeholders
            if key_var.get().strip() and val_var.get().strip()
        }

        if not data:
            messagebox.showwarning("Warning", "No placeholders to save.")
            return

        # Optional: check for duplicate keys
        keys = [key_var.get() for key_var, _, *_ in self.placeholders]
        if len(keys) != len(set(keys)):
            messagebox.showwarning("Warning", "Duplicate placeholder keys found.")
            return

        # Ask user for profile name
        profile_name = simpledialog.askstring("Save Profile", "Enter a name for this profile:")
        if not profile_name:
            return

        try:
            save_path = os.path.join(self.get_profiles_dir(), f"{profile_name}.json")
            with open(save_path, "w") as f:
                json.dump(data, f, indent=2)
            messagebox.showinfo("Saved", f"Profile '{profile_name}' saved.")
            self.refresh_profile_menu()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save profile:\n{e}")

    def load_named_profile(self):
        profile_name = self.profile_var.get()
        if not profile_name:
            return

        try:
            load_path = os.path.join(self.get_profiles_dir(), f"{profile_name}.json")
            with open(load_path, "r") as f:
                data = json.load(f)

            for widget in self.entries_frame.winfo_children():
                widget.destroy()
            self.placeholders.clear()
            self.add_table_headers()

            for key, val in data.items():
                self.add_placeholder_row()
                self.placeholders[-1][0].set(key)
                self.placeholders[-1][1].set(val)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load profile:\n{e}")

    def refresh_profile_menu(self):
        try:
            profiles = [f[:-5] for f in os.listdir(self.get_profiles_dir()) if f.endswith(".json")]
            self.profile_menu["values"] = profiles
            self.profile_var.set("Select Profile" if profiles else "No profiles found")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load profiles:\n{e}")
    
    def select_and_load_profile(self, profile_name):
        self.profile_var.set(profile_name)
        self.load_named_profile()
    
    def delete_selected_profile(self):
        profile_name = self.profile_var.get()
        if not profile_name or profile_name in ("Select Profile", "No profiles found"):
            messagebox.showwarning("No Profile Selected", "Please select a valid profile to delete.")
            return

        confirm = messagebox.askyesno("Delete Profile", f"Are you sure you want to delete '{profile_name}'?")
        if not confirm:
            return

        try:
            profile_path = os.path.join(self.get_profiles_dir(), f"{profile_name}.json")
            if os.path.exists(profile_path):
                os.remove(profile_path)
                messagebox.showinfo("Deleted", f"Profile '{profile_name}' has been deleted.")
                self.refresh_profile_menu()
            else:
                messagebox.showerror("Error", "Profile file not found.")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to delete profile:\n{e}")

    def extract_placeholders_from_template(self, path):
        placeholders = set()

        if path.endswith(".docx"):
            doc = Document(path)

            # Scan regular paragraphs
            for para in doc.paragraphs:
                matches = re.findall(r"\{(.*?)\}", para.text)
                placeholders.update(matches)

            # 🔧 Scan tables too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        matches = re.findall(r"\{(.*?)\}", cell.text)
                        placeholders.update(matches)

        elif path.endswith(".xlsx"):
            wb = load_workbook(path)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    for cell in row:
                        if isinstance(cell, str):
                            matches = re.findall(r"\{(.*?)\}", cell)
                            placeholders.update(matches)

        return sorted(placeholders)

    def handle_drop(self, event):
        file_path = event.data.strip().strip("{").strip("}")  # Handles Windows paths with spaces
        if file_path.lower().endswith((".docx", ".xlsx")) and os.path.exists(file_path):
            self.template_path.set(file_path)
            try:
                keys = self.extract_placeholders_from_template(file_path)
                if keys:
                    self.clear_all_fields()
                    for i, key in enumerate(keys):
                        if i < len(self.placeholders):
                            self.placeholders[i][0].set(key)
                            self.placeholders[i][1].set("")
                        else:
                            self.add_placeholder_row()
                            self.placeholders[-1][0].set(key)
                            self.placeholders[-1][1].set("")
                else:
                    messagebox.showinfo("No Placeholders Found", "No placeholders were detected in the template.")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to extract placeholders:\n{e}")
        else:
            messagebox.showwarning("Invalid File", "Please drop a valid .docx or .xlsx file.")

    def duplicate_placeholder_row(self, index):
        key_text = self.placeholders[index][0].get()
        value_text = self.placeholders[index][1].get()
        self.add_placeholder_row(key_text, value_text)

if __name__ == "__main__":
    root = TkinterDnD.Tk()

    # Optional: Set global font styles before launching the app
    default_font = tkFont.nametofont("TkDefaultFont")
    default_font.configure(size=12)
    tkFont.nametofont("TkTextFont").configure(size=12)
    tkFont.nametofont("TkFixedFont").configure(size=12)
    tkFont.nametofont("TkMenuFont").configure(size=12)
    tkFont.nametofont("TkHeadingFont").configure(size=18, weight="bold")

    # Launch your app
    app = TemplateFillerApp(root)
    root.mainloop()


